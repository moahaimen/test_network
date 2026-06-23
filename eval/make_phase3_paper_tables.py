#!/usr/bin/env python3
"""Build paper-ready Phase-3 tables (CSV + DOCX) from final outputs."""

from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


GEN_PATH = Path("results/phase3_final/GENERALIZATION_SUMMARY.csv")
FAIL_PATH = Path("results/phase3_final/FAILURE_SUMMARY.csv")
OUT_DIR = Path("results/phase3_final/paper_tables")


def _build_table_a(gen: pd.DataFrame) -> pd.DataFrame:
    if "tm_source" not in gen.columns:
        gen["tm_source"] = gen["source"].astype(str).str.lower().map(lambda s: "real" if s == "sndlib" else "mgm")

    ecmp = gen[gen["method"] == "ecmp"][["topology_id", "regime", "mean_mlu", "p95_mlu"]].rename(
        columns={"mean_mlu": "ecmp_mean_mlu", "p95_mlu": "ecmp_p95_mlu"}
    )

    best_idx = gen.groupby(["topology_id", "regime"])["mean_mlu"].idxmin()
    best = gen.loc[best_idx].copy()
    best = best.merge(ecmp, on=["topology_id", "regime"], how="left")
    best["gain_vs_ecmp_mean_pct"] = (
        (best["ecmp_mean_mlu"] - best["mean_mlu"]) / best["ecmp_mean_mlu"] * 100.0
    )

    if "mean_gap_pct" in best.columns:
        best["mean_gap_pct_vs_opt"] = best["mean_gap_pct"]
        best["mean_achieved_pct_vs_opt"] = best["mean_achieved_pct"]
    if "opt_solved_steps" in best.columns and "opt_total_steps" in best.columns:
        best["opt_solved_steps/opt_total_steps"] = best["opt_solved_steps"].fillna(0).astype(int).astype(str) + "/" + best[
            "opt_total_steps"
        ].fillna(0).astype(int).astype(str)

    cols = [
        "display_name",
        "source",
        "tm_source",
        "num_nodes",
        "num_edges",
        "regime",
        "method",
        "mean_mlu",
        "p95_mlu",
        "ecmp_mean_mlu",
        "ecmp_p95_mlu",
        "gain_vs_ecmp_mean_pct",
    ]
    if "mean_gap_pct_vs_opt" in best.columns:
        cols.append("mean_gap_pct_vs_opt")
    if "mean_achieved_pct_vs_opt" in best.columns:
        cols.append("mean_achieved_pct_vs_opt")
    if "opt_solved_steps/opt_total_steps" in best.columns:
        cols.append("opt_solved_steps/opt_total_steps")
    if "k_crit_used" in best.columns:
        cols.append("k_crit_used")

    return best[cols].sort_values(["regime", "source", "display_name"])


def _build_table_b(fail: pd.DataFrame) -> pd.DataFrame:
    required_cols = {
        "mean_mlu_reachable",
        "dropped_demand_pct",
        "unreachable_od_ratio",
        "feasible",
        "post_failure_peak_mlu",
        "recovery_steps",
    }
    missing = required_cols - set(fail.columns)
    if missing:
        raise RuntimeError(f"FAILURE_SUMMARY.csv is missing required columns: {sorted(missing)}")

    agg = {
        "mean_mlu_reachable": ("mean_mlu_reachable", "mean"),
        "p95_mlu_reachable": ("mean_mlu_reachable", lambda s: float(np.percentile(s, 95))),
        "peak_mlu_post_failure": ("post_failure_peak_mlu", "mean"),
        "dropped_demand_pct": ("dropped_demand_pct", "mean"),
        "unreachable_od_ratio": ("unreachable_od_ratio", "mean"),
        "feasible_runs_pct": ("feasible", lambda s: float(np.mean(s.astype(bool)) * 100.0)),
        "unrecovered_runs_pct": ("recovery_steps", lambda s: float(np.mean((s == -1).astype(float)) * 100.0)),
        "recovery_steps_recovered": (
            "recovery_steps",
            lambda s: float(s[s >= 0].mean()) if (s >= 0).any() else np.nan,
        ),
        "N": ("recovery_steps", "size"),
    }
    if "k_crit_used" in fail.columns:
        agg["k_crit_used"] = ("k_crit_used", "mean")

    return fail.groupby(["failure_type", "method"]).agg(**agg).reset_index().sort_values([
        "failure_type",
        "mean_mlu_reachable",
    ])


def _write_docx(table_a: pd.DataFrame, table_b: pd.DataFrame, out_path: Path) -> None:
    doc = Document()
    title = doc.add_paragraph("Phase-3 Paper Tables (Named Topologies)")
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    doc.add_paragraph("Table A - Generalization (Best vs ECMP)").runs[0].bold = True
    headers_a = [
        "Topology",
        "Source",
        "TM",
        "|V|",
        "|E|",
        "Regime",
        "Best method",
        "Best mean MLU",
        "Best p95 MLU",
        "ECMP mean MLU",
        "ECMP p95 MLU",
        "Gain vs ECMP (mean %)",
    ]
    include_gap_a = "mean_gap_pct_vs_opt" in table_a.columns
    include_ach_a = "mean_achieved_pct_vs_opt" in table_a.columns
    include_opt_steps_a = "opt_solved_steps/opt_total_steps" in table_a.columns
    include_kcrit_a = "k_crit_used" in table_a.columns
    if include_gap_a:
        headers_a.append("Mean gap vs LP-opt (%)")
    if include_ach_a:
        headers_a.append("Mean achieved vs LP-opt (%)")
    if include_opt_steps_a:
        headers_a.append("LP-opt solved/total")
    if include_kcrit_a:
        headers_a.append("Kcrit used")

    t_a = doc.add_table(rows=1, cols=len(headers_a))
    t_a.style = "Table Grid"
    for i, h in enumerate(headers_a):
        t_a.rows[0].cells[i].text = h

    for _, r in table_a.iterrows():
        row = t_a.add_row().cells
        row[0].text = str(r["display_name"])
        row[1].text = str(r["source"])
        row[2].text = str(r["tm_source"])
        row[3].text = str(int(r["num_nodes"]))
        row[4].text = str(int(r["num_edges"]))
        row[5].text = str(r["regime"])
        row[6].text = str(r["method"])
        row[7].text = f"{r['mean_mlu']:.4f}"
        row[8].text = f"{r['p95_mlu']:.4f}"
        row[9].text = f"{r['ecmp_mean_mlu']:.4f}"
        row[10].text = f"{r['ecmp_p95_mlu']:.4f}"
        col = 11
        row[col].text = f"{r['gain_vs_ecmp_mean_pct']:.2f}"
        if include_gap_a:
            col += 1
            row[col].text = "nan" if pd.isna(r["mean_gap_pct_vs_opt"]) else f"{r['mean_gap_pct_vs_opt']:.2f}"
        if include_ach_a:
            col += 1
            row[col].text = "nan" if pd.isna(r["mean_achieved_pct_vs_opt"]) else f"{r['mean_achieved_pct_vs_opt']:.2f}"
        if include_opt_steps_a:
            col += 1
            row[col].text = str(r["opt_solved_steps/opt_total_steps"])
        if include_kcrit_a:
            col += 1
            row[col].text = str(int(round(float(r["k_crit_used"]))))

    doc.add_paragraph("")
    doc.add_paragraph("Table B - Failure Scenarios (Reachable-only + Infeasibility)").runs[0].bold = True
    headers_b = [
        "Failure type",
        "Method",
        "Mean MLU (reachable)",
        "P95 MLU (reachable)",
        "Peak MLU",
        "Dropped demand %",
        "Unreachable OD %",
        "Feasible runs %",
        "Unrecovered runs %",
        "Avg recovery steps (recovered)",
        "N",
    ]
    include_kcrit_b = "k_crit_used" in table_b.columns
    if include_kcrit_b:
        headers_b.append("Kcrit used")

    t_b = doc.add_table(rows=1, cols=len(headers_b))
    t_b.style = "Table Grid"
    for i, h in enumerate(headers_b):
        t_b.rows[0].cells[i].text = h

    for _, r in table_b.iterrows():
        row = t_b.add_row().cells
        row[0].text = str(r["failure_type"])
        row[1].text = str(r["method"])
        row[2].text = f"{r['mean_mlu_reachable']:.4f}"
        row[3].text = f"{r['p95_mlu_reachable']:.4f}"
        row[4].text = f"{r['peak_mlu_post_failure']:.4f}"
        row[5].text = f"{r['dropped_demand_pct'] * 100.0:.2f}"
        row[6].text = f"{r['unreachable_od_ratio'] * 100.0:.3f}"
        row[7].text = f"{r['feasible_runs_pct']:.1f}"
        row[8].text = f"{r['unrecovered_runs_pct']:.1f}"
        row[9].text = "" if pd.isna(r["recovery_steps_recovered"]) else f"{r['recovery_steps_recovered']:.3f}"
        row[10].text = str(int(r["N"]))
        if include_kcrit_b:
            row[11].text = str(int(round(float(r["k_crit_used"]))))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def _write_report_md(table_a: pd.DataFrame, table_b: pd.DataFrame, out_path: Path) -> None:
    lines: list[str] = []
    lines.append("# Phase-3 Paper Tables Report")
    lines.append("")
    lines.append("## Table A (Best Method Per Topology/Regime)")
    lines.append("")
    lines.append(table_a.to_markdown(index=False))
    lines.append("")
    lines.append("## Table B (Failure Aggregates)")
    lines.append("")
    lines.append(table_b.to_markdown(index=False))
    lines.append("")
    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    if not GEN_PATH.exists():
        raise FileNotFoundError(f"Missing file: {GEN_PATH}")
    if not FAIL_PATH.exists():
        raise FileNotFoundError(f"Missing file: {FAIL_PATH}")

    OUT_DIR.mkdir(parents=True, exist_ok=True)

    gen = pd.read_csv(GEN_PATH)
    fail = pd.read_csv(FAIL_PATH)

    table_a = _build_table_a(gen)
    table_b = _build_table_b(fail)

    table_a.to_csv(OUT_DIR / "TableA_Generalization.csv", index=False)
    table_b.to_csv(OUT_DIR / "TableB_Failures.csv", index=False)

    _write_docx(table_a, table_b, OUT_DIR / "Phase3_Paper_Tables.docx")
    _write_report_md(table_a, table_b, OUT_DIR / "report.md")

    print(f"CSV: {OUT_DIR / 'TableA_Generalization.csv'}")
    print(f"CSV: {OUT_DIR / 'TableB_Failures.csv'}")
    print(f"DOCX: {OUT_DIR / 'Phase3_Paper_Tables.docx'}")
    print(f"MD: {OUT_DIR / 'report.md'}")


if __name__ == "__main__":
    main()
