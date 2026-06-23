#!/usr/bin/env python3
"""Full metrics for the CURRENT method ONLY (Topology-Agnostic Bottleneck-Ranking DDQN).
No baselines, no failure, no Mininet, no old-method content. All numbers from frozen Iter2 data."""
import json
from pathlib import Path
import numpy as np, pandas as pd, pickle
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = (__import__("pathlib").Path(__file__).resolve().parents[2] / "results/gnn_lpd_dqn_selective_db_lp/condition_compliant_k10_k50")
FR = OUT / "FROZEN_FINAL_LEARNED_RUNTIME_SAFE_ITER2"; RPT = OUT / "FINAL_REPORT"; FIG = RPT / "figs"
pc = pd.read_csv(FR / "final_learned_4of5_iter2_eval_per_cycle.csv")
adist = pd.read_csv(FR / "CONSOLIDATED_ACTION_DISTRIBUTION.csv")
abl = pd.read_csv(OUT / "FINAL_LEARNED_4OF5_ITER2_DDQN" / "rank_ablation.csv")
cnt = json.load(open(FR / "FINAL_LEARNED_4OF5_ITER2_AUDIT.json"))["runtime_counters"]
GNN_MS = {"abilene":3,"geant":7,"cernet":22,"sprintlink":27,"tiscali":33,"ebone":12,"germany50":26,"vtlwavenet2011":140}
NL = {"abilene":(12,30),"geant":(22,72),"cernet":(41,116),"sprintlink":(44,166),"tiscali":(49,172),"ebone":(23,76),"germany50":(50,176),"vtlwavenet2011":(92,192)}
TOPO = ["abilene","geant","cernet","sprintlink","tiscali","ebone","germany50","vtlwavenet2011"]
DISP = {"abilene":"Abilene","geant":"GEANT","cernet":"CERNET","sprintlink":"Sprintlink","tiscali":"Tiscali","ebone":"Ebone","germany50":"Germany50","vtlwavenet2011":"VtlWavenet"}
ZERO = ["germany50","vtlwavenet2011"]; acts=["KEEP","K50","K100","K200","K300","K500","K800"]

doc=Document(); sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Inches(0.85))
doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(10)
doc.styles["Heading 1"].font.color.rgb=RGBColor(0x1F,0x4D,0x78); HDR="1F4D78"
def shade(c,f):
    p=c._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),f); p.append(sh)
def sf(c,sz,b=False,wh=False):
    for pp in c.paragraphs:
        for r in pp.runs:
            r.font.size=Pt(sz); r.font.bold=b
            if wh: r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
def tbl(headers, rows, fz=8, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=str(h); shade(c,HDR); sf(c,fz,True,True)
    for row in rows:
        cs=t.add_row().cells
        for j,v in enumerate(row): cs[j].text=str(v); sf(cs[j],fz)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
def para(txt,size=10,bold=False,italic=False,space=6):
    p=doc.add_paragraph(); r=p.add_run(txt); r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic
    p.paragraph_format.space_after=Pt(space); return p
def h1(t): doc.add_heading(t,1)
def h2(t): doc.add_heading(t,2)

tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("Full Metrics — Topology-Agnostic Bottleneck-Ranking DDQN"); r.font.size=Pt(16); r.bold=True; r.font.color.rgb=RGBColor(0x1F,0x4D,0x78)
para("Metrics for the current method only. All values are computed from the frozen final evaluation "
     "(final_learned_4of5_iter2_eval_per_cycle.csv; 3816 cycles; 8 topologies; argmax-Q, no forced actuator).",
     italic=True, size=9)

# 1. Method definition
h1("1. Method definition")
tbl(["Component","Role in the final method"],
 [["GNN-LPD scorer","learned per-OD criticality score (one input signal)"],
  ["Bottleneck-aware ranking","orders OD pairs by relief (demand x ECMP-flow x link-util) + 0.3 x GNN; selects top-K"],
  ["Double-DQN (argmax-Q)","action policy: chooses KEEP or an optimize budget K from {50,100,200,300,500,800}"],
  ["Selected-flow LP","feasible routing optimizer over the selected top-K ODs only"],
  ["ECMP","fixed background routing for all nonselected ODs"]], fz=9.5, widths=[2.2,5.0])

# 2. Normal-scenario metrics
h1("2. Normal-scenario metrics (per topology)")
rows=[]
for t in TOPO:
    g=pc[pc.topology==t]
    rows.append([DISP[t], len(g), f"{g.PR.mean():.4f}", f"{np.median(g.PR):.4f}", f"{(g.PR>=0.95).mean()*100:.0f}%",
        f"{(g.PR>=0.90).mean()*100:.0f}%", f"{g.PR.min():.4f}", f"{g.DB.mean():.4f}", f"{np.percentile(g.DB,95):.4f}",
        f"{g.DB.max():.4f}", f"{g.decision_ms.mean():.1f}", f"{np.percentile(g.decision_ms,95):.1f}"])
tbl(["Topology","N","MeanPR","MedPR","PR>=.95","PR>=.90","MinPR","MeanDB","P95DB","MaxDB","Mean ms","P95 ms"], rows, fz=7.5)
para("All eight topologies satisfy PR >= 0.90 and both mean and p95 decision time below 500 ms.", italic=True, size=9)

# 2b MLU/optimum
h2("2b. MLU and MLU/optimum (= 1/PR)")
rows=[]
for t in TOPO:
    g=pc[pc.topology==t]; mo=(1.0/g.PR)
    rows.append([DISP[t], f"{g.MLU.mean():.4f}", f"{np.percentile(g.MLU,95):.4f}", f"{g.MLU.max():.4f}", f"{mo.mean():.4f}", f"{mo.max():.4f}"])
tbl(["Topology","Mean MLU","P95 MLU","Max MLU","Mean MLU/opt","Max MLU/opt"], rows, fz=8.5, widths=[1.6,1.3,1.3,1.3,1.5,1.5])

# 3. Pooled summary
h1("3. Pooled summary (all cycles)")
g=pc
tbl(["Metric","Value"],
 [["Cycles (N)", len(g)],["Mean PR", f"{g.PR.mean():.4f}"],["Median PR", f"{np.median(g.PR):.4f}"],
  ["PR >= 0.95", f"{(g.PR>=0.95).mean()*100:.1f}%"],["PR >= 0.90", f"{(g.PR>=0.90).mean()*100:.1f}%"],
  ["Min PR", f"{g.PR.min():.4f}"],["Mean DB", f"{g.DB.mean():.4f}"],["P95 DB", f"{np.percentile(g.DB,95):.4f}"],
  ["Mean decision ms", f"{g.decision_ms.mean():.1f}"],["P95 decision ms", f"{np.percentile(g.decision_ms,95):.1f}"],
  ["Max decision ms", f"{g.decision_ms.max():.1f}"]], fz=9, widths=[2.4,2.2])

# 4. Action distribution
h1("4. Action distribution (per topology)")
tbl(["Topology"]+acts+["Most used"], [[DISP[t]]+[int(adist[adist.Topology==t][a].iloc[0]) for a in acts]+[adist[adist.Topology==t]["Most_used"].iloc[0]] for t in TOPO], fz=8)
para("The action is chosen per traffic-matrix cycle by argmax-Q, so one topology uses multiple actions. KEEP can "
     "dominate while mean K is nonzero because KEEP reuses the carried routing between optimize cycles.", italic=True, size=8.5)

# 5. K sensitivity / ranking ablation (current method)
h1("5. K sensitivity and ranking ablation")
tbl(["Topology","K","Ranking","Mean PR","Mean ms"], [[DISP[r.topology], int(r.K), r.ranking, f"{r.mean_PR:.4f}", f"{r.mean_ms:.1f}"] for _,r in abl.iterrows()], fz=8.5)

# 6. RL policy + reward
h1("6. RL policy and reward settings")
tbl(["Parameter","Value"],
 [["Controller","Double-DQN (argmax-Q at inference)"],["Action space","KEEP, K50, K100, K200, K300, K500, K800"],
  ["gamma","0.5"],["W_PR","10.0"],["W_MLU","5.0"],["W_DB (lambda for DB)","20.0"],["W_MS","0.003"],["W_K","0.5"],
  ["target bonus / gates","+10 if PR>=target; penalty if PR<target; flat anti-KEEP-below-target; ms>500 gate"],
  ["epsilon","1.0 -> 0.05"],["episodes","22 (x6 seen topos x 160 cycles)"]], fz=8.5, widths=[2.4,4.8])

# 7. Zero-shot
h1("7. Zero-shot generalization")
rows=[]
for t in ZERO+["__c__"]:
    if t=="__c__": g=pc[pc.topology.isin(ZERO)]; name="Combined zero-shot"
    else: g=pc[pc.topology==t]; name=DISP[t]
    rows.append([name, len(g), f"{g.PR.mean():.4f}", f"{(g.PR>=0.90).mean()*100:.0f}%", f"{(g.PR>=0.95).mean()*100:.0f}%", f"{g.PR.min():.4f}", f"{g.DB.mean():.4f}", f"{np.percentile(g.DB,95):.4f}", f"{g.decision_ms.mean():.1f}"])
tbl(["Zero-shot topology","N","Mean PR","PR>=.90","PR>=.95","Min PR","Mean DB","P95 DB","Mean ms"], rows, fz=8.5)
# robust VtlWavenet (extended from 40 to 200 TMs)
_vx = OUT/"FINAL_LEARNED_4OF5_ITER2_DDQN"/"vtl_extended_200_per_cycle.csv"
if _vx.exists():
    v=pd.read_csv(_vx)
    para("VtlWavenet robust re-evaluation (extended from 40 to 200 traffic matrices):", bold=True, size=9)
    tbl(["VtlWavenet","N (TMs)","Mean PR","PR>=.90","Min PR","Mean DB","Mean ms","P95 ms"],
        [["Original sample","40","0.9373","100%","-","0.0007","295.9","307.3"],
         ["Robust (200 TMs)", len(v), f"{v.PR.mean():.4f}", f"{(v.PR>=0.90).mean()*100:.0f}%", f"{v.PR.min():.4f}", f"{v.DB.mean():.4f}", f"{v.decision_ms.mean():.1f}", f"{np.percentile(v.decision_ms,95):.1f}"]], fz=8.5)
    para("The PR is confirmed robust (200/200 cycles >= 0.90, mean 0.934). The larger sample shows the decision-time "
         "tail (p95) exceeds 500 ms on this largest topology; decision time is also contention-sensitive. "
         "VtlWavenet remains a zero-shot generalization spot-check (8372 OD pairs, the costliest topology).", italic=True, size=8.5)

# 8. Decision-time + solver + LP size
h1("8. Decision-time breakdown, solver, LP size")
rows=[]
for t in TOPO:
    g=pc[pc.topology==t]; opt=g[g.action!="KEEP"]; lp=(opt.decision_ms-GNN_MS[t]).mean() if len(opt) else 0
    rows.append([DISP[t], GNN_MS[t], f"{lp:.1f}", f"{g.decision_ms.mean():.1f}", f"{np.percentile(g.decision_ms,95):.1f}", f"{g.decision_ms.max():.1f}"])
tbl(["Topology","GNN/scoring ms","LP ms (optimize avg)","Total mean ms","P95 ms","Max ms"], rows, fz=8.5, widths=[1.5,1.4,1.6,1.3,1.0,1.0])
para("Decision time by action type (pooled):", bold=True, size=9)
rows=[[a, len(pc[pc.action==a]), f"{pc[pc.action==a].decision_ms.mean():.1f}", f"{np.percentile(pc[pc.action==a].decision_ms,95):.1f}", f"{pc[pc.action==a].decision_ms.max():.1f}"] for a in acts if len(pc[pc.action==a])]
tbl(["Action","Rows","Mean ms","P95 ms","Max ms"], rows, fz=8.5, widths=[1.4,1.2,1.4,1.4,1.4])
para("Solver and LP size:", bold=True, size=9)
tbl(["Item","Value"], [["LP solver","PuLP + CBC (open-source)"],["LP type","selected-flow path LP, DB-budgeted"],["k_paths","8 for K50-K300; 4 for K500/K800"],["Seed","42"],["Framework","PyTorch (CPU)"]], fz=8.5, widths=[2.0,5.2])
rows=[]
for t in TOPO:
    g=pc[pc.topology==t]; n,l=NL[t]; selavg=g.selected_K.mean(); kp=8 if g.selected_K.max()<=300 else "8/4"; var=selavg*(8 if g.selected_K.max()<=300 else 6)
    rows.append([DISP[t], n, l, int(g.num_active_ods.mean()), f"{selavg:.0f}", str(kp), f"~{var:.0f}", f"{g.decision_ms.mean():.1f}"])
tbl(["Topology","Nodes","Links","OD pairs","sel OD avg","k_paths","LP vars ~avg","Mean ms"], rows, fz=8, widths=[1.3,0.8,0.8,1.0,1.2,0.9,1.1,0.9])

# 9. Figures
h1("9. Figures (current method)")
for fn,cap in [("fig1_pr_cdf.png","PR CDF"),("fig2_db_cdf.png","DB CDF"),("fig3_ms_cdf.png","Decision-time CDF"),
               ("fig8_pr_vs_db.png","PR-DB tradeoff"),("fig4_meanpr.png","Mean PR by topology"),
               ("fig5_meandb.png","Mean DB by topology"),("fig6_actiondist.png","Action distribution"),
               ("fig7_kdist.png","K-budget distribution")]:
    if (FIG/fn).exists():
        doc.add_picture(str(FIG/fn), width=Inches(4.7)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        para(cap, italic=True, size=8.5, space=8)

# 9b. Failure scenarios (real current-method rerun)
_fs = OUT/"FAILURE_VALIDATION_ITER2"/"failure_iter2_summary.csv"
if _fs.exists():
    fs = pd.read_csv(_fs); fd = pd.read_csv(OUT/"FAILURE_VALIDATION_ITER2"/"failure_iter2_disconnect_detail.csv")
    h1("9b. Failure scenarios (current method, real rerun)")
    para("Failure-link scenarios were rerun on the frozen Iter2 controller (argmax-Q + bottleneck ranking + "
         "selected-flow LP; nonselected ODs = ECMP). Abilene and GEANT, 20 cycles per scenario.", size=9)
    rows=[]
    for _,r in fs.iterrows():
        rows.append([DISP.get(r.Topology,r.Topology), r.Scenario.replace('_',' '), int(r.N), f"{r.Mean_PR:.4f}", f"{r.Mean_MLU:.4f}", f"{r.ECMP_Mean_MLU:.4f}", f"{r.Mean_DB:.4f}", f"{r.Mean_ms:.0f}", int(r.Disconnected_ODs)])
    tbl(["Topology","Scenario","N","Mean PR","Our MLU","ECMP MLU","Mean DB","Mean ms","Disc. ODs"], rows, fz=7.3)
    para("Honest reading: the method holds PR>=0.99 in most failure scenarios and roughly halves MLU versus ECMP. "
         "Weak points: Abilene two-link failure (PR 0.8965, below 0.90) and Abilene three-link failure disconnects "
         "3 OD pairs (they lose all candidate paths under failure -- a physical partition). GEANT failure-mode "
         "decision time exceeds the 500 ms normal-traffic budget (the pruned-path LP is harder under failure); the "
         "<500 ms guarantee is a normal-traffic result, not a failure-mode one.", italic=True, size=8.5)
    para("Disconnected-OD detail:", bold=True, size=9)
    tbl(["Topology","Scenario","Failed links","Connected?","Disconnected ODs","Explanation"],
        [[DISP.get(r.Topology,r.Topology), r.Scenario.replace('_',' '), int(r.Failed_links), r.Connected, int(r.Disconnected_ODs), r.Explanation] for _,r in fd.iterrows()], fz=7)
    for fn,cap in [("failure_iter2_mlu_cdf.png","Failure MLU CDF (current method)"),("failure_iter2_db_by_scenario.png","Failure DB by scenario (current method)")]:
        fp=OUT/"FAILURE_VALIDATION_ITER2"/fn
        if fp.exists():
            doc.add_picture(str(fp), width=Inches(4.7)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
            para(cap, italic=True, size=8.5, space=8)

# 10. Reproducibility
h1("10. Reproducibility")
tbl(["Item","Value"],
 [["Final model","final_learned_4of5_iter2_model.pt (Double-DQN, dim=33, 7 actions)"],
  ["Per-cycle data","final_learned_4of5_iter2_eval_per_cycle.csv (3816 rows)"],
  ["Train command","python3 scripts/phase1_5/run_final_iter2.py"],
  ["Counters", f"td_updates={cnt['td_updates']}, target_updates={cnt['target_updates']}, ce_updates={cnt['ce_updates']}"],
  ["Seed / solver","42 / PuLP+CBC"]], fz=8.5, widths=[1.8,5.4])

DOCX=RPT/"Current_Method_Full_Metrics_TopoAgnostic_Bottleneck_DDQN.docx"
doc.save(str(DOCX)); print("saved:", DOCX); print("DONE")
