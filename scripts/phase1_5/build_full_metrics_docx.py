#!/usr/bin/env python3
"""Full-metrics DOCX for the FINAL method (Topology-Agnostic Bottleneck-Ranking DDQN).
Every metric computed from the frozen Iter2 per-cycle CSV + prepass. Items not rerun for the
final controller (failure scenarios, live Mininet, non-ECMP baselines) are marked honestly."""
import json
from pathlib import Path
import numpy as np, pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = (__import__("pathlib").Path(__file__).resolve().parents[2] / "results/gnn_lpd_dqn_selective_db_lp/condition_compliant_k10_k50")
FR = OUT / "FROZEN_FINAL_LEARNED_RUNTIME_SAFE_ITER2"; RPT = OUT / "FINAL_REPORT"; FIG = RPT / "figs"
pc = pd.read_csv(FR / "final_learned_4of5_iter2_eval_per_cycle.csv")
adist = pd.read_csv(FR / "CONSOLIDATED_ACTION_DISTRIBUTION.csv")
import pickle; P = pickle.load(open(OUT / "_prepass.pkl", "rb"))
abl = pd.read_csv(OUT / "FINAL_LEARNED_4OF5_ITER2_DDQN" / "rank_ablation.csv")
cnt = json.load(open(FR / "FINAL_LEARNED_4OF5_ITER2_AUDIT.json"))["runtime_counters"]
W = {"abilene":(2016,4032),"geant":(672,1344),"cernet":(200,400),"sprintlink":(200,400),"tiscali":(200,400),"ebone":(200,400),"germany50":(0,288),"vtlwavenet2011":(0,40)}
GNN_MS = {"abilene":3,"geant":7,"cernet":22,"sprintlink":27,"tiscali":33,"ebone":12,"germany50":26,"vtlwavenet2011":140}
NL = {"abilene":(12,30),"geant":(22,72),"cernet":(41,116),"sprintlink":(44,166),"tiscali":(49,172),"ebone":(23,76),"germany50":(50,176),"vtlwavenet2011":(92,192)}
TOPO = list(W); DISP = {"abilene":"Abilene","geant":"GEANT","cernet":"CERNET","sprintlink":"Sprintlink","tiscali":"Tiscali","ebone":"Ebone","germany50":"Germany50","vtlwavenet2011":"VtlWavenet"}
SEEN = ["abilene","geant","cernet","sprintlink","tiscali","ebone"]; ZERO = ["germany50","vtlwavenet2011"]
def ecmp_pr(t):
    lo,hi=W[t]; d=P[(t,lo,hi)]; return float(np.mean([min(1,d['opt'][x]/d['emlu'][x]) if d['emlu'][x]>0 else 0 for x in range(lo,hi)]))

doc=Document(); sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Inches(0.85))
doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(10)
doc.styles["Heading 1"].font.color.rgb=RGBColor(0x1F,0x4D,0x78); HDR="1F4D78"
def shade(c,f):
    p=c._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),f); p.append(sh)
def sf(c,sz,b=False,wh=False):
    for pp in c.paragraphs:
        for r in pp.runs:
            r.font.size=Pt(sz); r.font.bold=b
            if wh: r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
def tbl(headers, rows, fz=8, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=str(h); shade(c,HDR); sf(c,fz,True,True)
    for row in rows:
        cs=t.add_row().cells
        for j,v in enumerate(row): cs[j].text=str(v); sf(cs[j],fz)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
def para(txt,size=10,bold=False,italic=False,space=6):
    p=doc.add_paragraph(); r=p.add_run(txt); r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic
    p.paragraph_format.space_after=Pt(space); return p
def h1(t): doc.add_heading(t,1)
def h2(t): doc.add_heading(t,2)

tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("Full Metrics — Topology-Agnostic Bottleneck-Ranking DDQN (final controller)"); r.font.size=Pt(16); r.bold=True; r.font.color.rgb=RGBColor(0x1F,0x4D,0x78)
para("All metrics below are computed from the frozen final evaluation (final_learned_4of5_iter2_eval_per_cycle.csv, "
     "3816 cycles, 8 topologies). Note: this final method is the Topology-Agnostic Bottleneck-Ranking DDQN "
     "(real Double-DQN, argmax-Q, bottleneck-aware OD ranking, selected-flow LP, nonselected ODs = ECMP). It is "
     "NOT the older Reward-Gated GNN-LPD method and uses no RandomForest and no reward gate; it is labelled "
     "accordingly rather than as 'RG-GNN-LPD'.", italic=True, size=9)

# 1. Method definition
h1("1. Method definition")
tbl(["Component","Role in the final method"],
 [["GNN-LPD scorer","learned per-OD criticality score (one input signal)"],
  ["Bottleneck-aware ranking","orders OD pairs by relief (demand x ECMP-flow x link-util) + 0.3 x GNN; selects top-K"],
  ["Double-DQN (argmax-Q)","action policy: chooses KEEP or an optimize budget K from {50,100,200,300,500,800}"],
  ["Selected-flow LP","feasible routing optimizer over the selected top-K ODs only"],
  ["ECMP","fixed background routing for all nonselected ODs"]], fz=9, widths=[2.2,5.0])

# 2. Dataset / split
h1("2. Dataset / train / validation / test")
tbl(["Topology","Source","Role","Train cycles","Eval cycles","Zero-shot?"],
 [["Abilene","SNDlib (real)","Seen","0-160 (of 0-2016)","2016-4032","No"],
  ["GEANT","SNDlib (real)","Seen","0-160","672-1344","No"],
  ["CERNET","real","Seen","0-160","200-400","No"],
  ["Sprintlink","Rocketfuel (synthetic caps)","Seen","0-160","200-400","No"],
  ["Tiscali","Rocketfuel (synthetic caps)","Seen","0-160","200-400","No"],
  ["Ebone","Rocketfuel (synthetic caps)","Seen","0-160","200-400","No"],
  ["Germany50","SNDlib (real)","Generalization","- (never trained)","0-288","Yes"],
  ["VtlWavenet","TopologyZoo (synthetic caps)","Generalization","- (never trained)","0-40","Yes"]], fz=8, widths=[1.3,1.9,1.3,1.5,1.2,0.9])
para("Note: CERNET, Sprintlink, Tiscali, Ebone, VtlWavenet use synthetic (degree-based) link capacities in this "
     "pipeline; Abilene, GEANT, Germany50 use real capacities. OD-fraction comparisons across papers are only "
     "apples-to-apples on the same capacity model.", italic=True, size=8.5)

# 3. Normal-scenario per-topology metrics
h1("3. Normal-scenario metrics (per topology)")
rows=[]
for t in TOPO:
    g=pc[pc.topology==t]
    rows.append([DISP[t], len(g), f"{g.PR.mean():.4f}", f"{np.median(g.PR):.4f}", f"{(g.PR>=0.95).mean()*100:.0f}%",
        f"{(g.PR>=0.90).mean()*100:.0f}%", f"{g.PR.min():.4f}", f"{g.DB.mean():.4f}", f"{np.percentile(g.DB,95):.4f}",
        f"{g.DB.max():.4f}", f"{g.decision_ms.mean():.1f}", f"{np.percentile(g.decision_ms,95):.1f}"])
tbl(["Topology","N","MeanPR","MedPR","PR>=.95","PR>=.90","MinPR","MeanDB","P95DB","MaxDB","Mean ms","P95 ms"], rows, fz=7.5)

# 3b MLU / optimum
h2("3b. MLU and MLU/optimum")
rows=[]
for t in TOPO:
    g=pc[pc.topology==t]; mo=(1.0/g.PR); rows.append([DISP[t], f"{g.MLU.mean():.4f}", f"{np.percentile(g.MLU,95):.4f}", f"{g.MLU.max():.4f}", f"{mo.mean():.4f}", f"{mo.max():.4f}"])
tbl(["Topology","Mean MLU","P95 MLU","Max MLU","Mean MLU/opt","Max MLU/opt"], rows, fz=8.5, widths=[1.6,1.3,1.3,1.3,1.5,1.5])
para("MLU/optimum = our MLU / strict-optimum MLU = 1/PR. Values near 1.0 mean near-optimal routing.", italic=True, size=8.5)

# 4. Baseline comparison (ECMP computable; others not rerun)
h1("4. Baseline comparison (same dataset/protocol)")
rows=[["ECMP (baseline)"]+[f"{np.mean([ecmp_pr(t) for t in TOPO]):.4f}", "computed from prepass (opt/ECMP-MLU) per topology"]]
tbl(["Method","Pooled Mean PR","Status"],
 [["ECMP (baseline)", f"{np.mean([ecmp_pr(t) for t in TOPO]):.4f}", "computed (see per-topology below)"],
  ["Final DDQN (ours)", f"{pc.PR.mean():.4f}", "computed (frozen Iter2)"],
  ["OSPF / shortest-path", "-", "NOT rerun for the final method/protocol"],
  ["Top-K by demand", "-", "NOT rerun for the final method/protocol"],
  ["Bottleneck / sensitivity selectors", "-", "NOT rerun for the final method/protocol"]], fz=8.5, widths=[2.6,1.4,3.2])
para("Per-topology ECMP vs final DDQN PR:", bold=True, size=9)
tbl(["Topology","ECMP PR","Final DDQN PR","Improvement"],
 [[DISP[t], f"{ecmp_pr(t):.4f}", f"{pc[pc.topology==t].PR.mean():.4f}", f"+{pc[pc.topology==t].PR.mean()-ecmp_pr(t):.4f}"] for t in TOPO], fz=8.5, widths=[1.8,1.6,1.8,1.6])
para("Only ECMP is recomputed on the exact final protocol; OSPF/top-K/bottleneck/sensitivity baselines were not "
     "rerun for the final controller and are therefore not claimed.", italic=True, size=8.5)

# 5. K sensitivity + action distribution
h1("5. K sensitivity and action distribution")
para("Ranking/K ablation (fixed K, carry-forward) — gate(blend) vs relief-only vs GNN-only:", bold=True, size=9)
tbl(["Topology","K","Ranking","Mean PR","Mean ms"], [[DISP[r.topology], int(r.K), r.ranking, f"{r.mean_PR:.4f}", f"{r.mean_ms:.1f}"] for _,r in abl.iterrows()], fz=8)
para("Action distribution by topology (per-cycle argmax-Q choices):", bold=True, size=9)
acts=["KEEP","K50","K100","K200","K300","K500","K800"]
tbl(["Topology"]+acts+["Most used"], [[DISP[t]]+[int(adist[adist.Topology==t][a].iloc[0]) for a in acts]+[adist[adist.Topology==t]["Most_used"].iloc[0]] for t in TOPO], fz=8)
para("One topology uses multiple actions because the action is chosen per traffic-matrix cycle, not once per "
     "topology; different TMs need different K. KEEP can dominate while mean K is nonzero because KEEP reuses the "
     "carried routing between optimize cycles.", italic=True, size=8.5)

# 6. Gate / RL policy + reward settings
h1("6. RL policy and reward settings")
tbl(["Parameter","Value"],
 [["Controller","Double-DQN (argmax-Q at inference; not a reward gate / not RandomForest)"],
  ["Action space","KEEP, K50, K100, K200, K300, K500, K800"],
  ["gamma (discount)","0.5"],["W_PR (PR reward)","10.0"],["W_MLU","5.0"],["W_DB (lambda for DB)","20.0"],
  ["W_MS (decision-time)","0.003"],["W_K (K/active penalty)","0.5"],
  ["target bonus / gate","+10 if PR>=target; strong penalty if PR<target; flat anti-KEEP-below-target; ms>500 gate"],
  ["epsilon","1.0 -> 0.05 over training"],["episodes","22 (x6 seen topos x 160 cycles)"]], fz=8.5, widths=[2.4,4.8])

# 7. Zero-shot
h1("7. Zero-shot generalization")
rows=[]
for t in ZERO+["__combined__"]:
    if t=="__combined__": g=pc[pc.topology.isin(ZERO)]; name="Combined zero-shot"
    else: g=pc[pc.topology==t]; name=DISP[t]
    rows.append([name, len(g), f"{g.PR.mean():.4f}", f"{(g.PR>=0.90).mean()*100:.0f}%", f"{(g.PR>=0.95).mean()*100:.0f}%", f"{g.PR.min():.4f}", f"{g.DB.mean():.4f}", f"{np.percentile(g.DB,95):.4f}", f"{g.decision_ms.mean():.1f}"])
tbl(["Zero-shot topology","N","Mean PR","PR>=.90","PR>=.95","Min PR","Mean DB","P95 DB","Mean ms"], rows, fz=8.5)
para("Germany50 and VtlWavenet were never seen in training; both achieve PR>=0.90 with mean decision time <500 ms.", italic=True, size=8.5)

# 8. Decision-time breakdown + solver + LP size
h1("8. Decision-time breakdown, solver, and LP size")
para("Decision-time components (GNN/DDQN scoring is a per-topology constant; LP is the remainder of optimize cycles):", bold=True, size=9)
rows=[]
for t in TOPO:
    g=pc[pc.topology==t]; opt=g[g.action!="KEEP"]; lp=(opt.decision_ms-GNN_MS[t]).mean() if len(opt) else 0
    rows.append([DISP[t], GNN_MS[t], f"{lp:.1f}", f"{g.decision_ms.mean():.1f}", f"{np.percentile(g.decision_ms,95):.1f}", f"{g.decision_ms.max():.1f}"])
tbl(["Topology","GNN/scoring ms","LP ms (optimize avg)","Total mean ms","P95 ms","Max ms"], rows, fz=8.5, widths=[1.5,1.4,1.6,1.3,1.0,1.0])
para("Decision time by action type (pooled over all topologies):", bold=True, size=9)
rows=[]
for a in acts:
    g=pc[pc.action==a]
    if len(g): rows.append([a, len(g), f"{g.decision_ms.mean():.1f}", f"{np.percentile(g.decision_ms,95):.1f}", f"{g.decision_ms.max():.1f}"])
tbl(["Action","Rows","Mean ms","P95 ms","Max ms"], rows, fz=8.5, widths=[1.4,1.2,1.4,1.4,1.4])
para("Solver / reproducibility:", bold=True, size=9)
tbl(["Item","Value"],
 [["LP solver","PuLP + CBC (open-source)"],["LP type","selected-flow path LP, DB-budgeted"],
  ["k_paths","8 for K50-K300; 4 for K500/K800"],["Seed","42"],["Framework","PyTorch (CPU)"]], fz=8.5, widths=[2.0,5.2])
para("LP problem size (per topology):", bold=True, size=9)
rows=[]
for t in TOPO:
    g=pc[pc.topology==t]; n,l=NL[t]; selavg=g.selected_K.mean(); kp=8 if g.selected_K.max()<=300 else "8/4"
    var = selavg* (8 if g.selected_K.max()<=300 else 6)
    rows.append([DISP[t], n, l, int(g.num_active_ods.mean()), f"{selavg:.0f}", str(kp), f"~{var:.0f}", f"{g.decision_ms.mean():.1f}"])
tbl(["Topology","Nodes","Links","OD pairs","selected OD avg","k_paths","LP vars ~avg","Mean ms"], rows, fz=8, widths=[1.3,0.8,0.8,1.0,1.3,0.9,1.1,0.9])

# 9. Failure (NOT rerun)
h1("9. Failure scenarios")
para("Failure-link scenarios (single/two/three-link failure, random link failure 1/2, capacity degradation 50%, "
     "mixed spike+failure) were NOT rerun for the frozen final controller. No failure metrics, disconnected-OD "
     "tables, or failure CDFs are claimed for the final method. The earlier offline failure artifacts in the "
     "repository belong to an earlier method and are not attributed to this controller.", bold=True, size=9.5)

# 10. Mininet / SDN (retained, separate)
h1("10. Mininet / SDN validation")
para("Live SDN/Mininet metrics are retained from the prior Phase 1.5 operational-validation artifact (Abilene and "
     "GEANT, normal + failure scenarios: throughput, RTT, jitter, loss, rule count, install ms, recovery ms, "
     "disconnected). They are separate evidence and the final Iter2 DDQN was NOT rerun inside Mininet. See the "
     "main report (Section 11) for the retained table. SDN rule-installation time is reported separately from "
     "offline decision time and must not be mixed with it.", bold=True, size=9.5)

# 11. Figures
h1("11. Figures (regenerated from final CSVs)")
for fn,cap in [("fig1_pr_cdf.png","PR CDF"),("fig2_db_cdf.png","DB CDF"),("fig3_ms_cdf.png","Decision-time CDF"),
               ("fig8_pr_vs_db.png","PR-DB tradeoff"),("fig4_meanpr.png","Mean PR by topology"),
               ("fig5_meandb.png","Mean DB by topology"),("fig6_actiondist.png","Action distribution"),
               ("fig7_kdist.png","K-budget distribution")]:
    if (FIG/fn).exists():
        doc.add_picture(str(FIG/fn), width=Inches(4.7)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        para(cap, italic=True, size=8.5, space=8)
para("Failure normalized-MLU CDF, failure DB-by-scenario, and the Mininet throughput/loss chart are NOT "
     "regenerated for the final method because failure scenarios were not rerun and the final controller was not "
     "rerun in Mininet.", italic=True, size=8.5)

# 12. Reproducibility
h1("12. Reproducibility")
tbl(["Item","Value"],
 [["Final model","final_learned_4of5_iter2_model.pt (Double-DQN, dim=33, 7 actions)"],
  ["Per-cycle data","final_learned_4of5_iter2_eval_per_cycle.csv (3816 rows)"],
  ["Train command","python3 scripts/phase1_5/run_final_iter2.py"],
  ["Verify command","python3 verify_results.py (in the reproduction package)"],
  ["Counters", f"td_updates={cnt['td_updates']}, target_updates={cnt['target_updates']}, ce_updates={cnt['ce_updates']}"],
  ["Seed / solver","42 / PuLP+CBC"]], fz=8.5, widths=[1.8,5.4])

DOCX=RPT/"Full_Metrics_TopoAgnostic_Bottleneck_DDQN.docx"
doc.save(str(DOCX)); print("saved:", DOCX); print("DONE")
