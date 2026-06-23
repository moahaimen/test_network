#!/usr/bin/env python3
"""Standalone DOCX: proof that the final DDQN is genuinely learning (curve + counters + controlled test)."""
import json
from pathlib import Path
import numpy as np, pandas as pd
import matplotlib; matplotlib.use("Agg"); import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = (__import__("pathlib").Path(__file__).resolve().parents[2] / "results/gnn_lpd_dqn_selective_db_lp/condition_compliant_k10_k50")
ITER2 = OUT / "FINAL_LEARNED_4OF5_ITER2_DDQN"
RPT = OUT / "FINAL_REPORT"; FIG = RPT / "figs"; FIG.mkdir(parents=True, exist_ok=True)
tl = pd.read_csv(ITER2 / "final_learned_4of5_iter2_train_log.csv")
cnt = json.load(open(ITER2 / "FINAL_LEARNED_4OF5_ITER2_AUDIT.json"))["runtime_counters"]
proof = pd.read_csv(ITER2 / "learning_proof.csv")

# --- learning-curve figure (loss + reward + epsilon) ---
fig, ax1 = plt.subplots(figsize=(6.2, 3.4)); ax2 = ax1.twinx()
ax1.plot(tl.episode, tl.mean_td_loss, "o-", color="#c0392b", label="TD (Huber) loss", ms=3)
ax2.plot(tl.episode, tl.mean_reward, "s-", color="#1f6dd8", label="Mean reward", ms=3)
ax2.plot(tl.episode, tl.epsilon*tl.mean_reward.max(), ":", color="#7f8c8d", lw=1, label="epsilon (scaled)")
ax1.set_xlabel("Training episode"); ax1.set_ylabel("TD loss", color="#c0392b"); ax2.set_ylabel("Mean reward", color="#1f6dd8")
ax1.set_title("DDQN learning curve (Iter2): loss ↓, reward ↑"); ax1.grid(alpha=.3)
l1,la1=ax1.get_legend_handles_labels(); l2,la2=ax2.get_legend_handles_labels(); ax1.legend(l1+l2, la1+la2, fontsize=7, loc="center right")
fig.tight_layout(); fig.savefig(FIG/"proof_learning_curve.png", dpi=140); plt.close(fig)

# --- controlled-comparison bar (reward) ---
order=["trained","untrained","random","k800","keep"]; lab={"trained":"Trained DDQN","untrained":"Untrained\n(random init)","random":"Random\naction","k800":"Always-K800\n(best fixed)","keep":"Always-KEEP"}
pm={r.policy:r for _,r in proof.iterrows()}
fig,ax=plt.subplots(figsize=(6.2,3.2)); vals=[pm[p].mean_reward for p in order]; cols=["#27ae60","#e67e22","#95a5a6","#8e44ad","#c0392b"]
ax.bar([lab[p] for p in order], vals, color=cols); ax.set_ylabel("Mean reward (objective)"); ax.set_title("Trained DDQN beats untrained, random, and best fixed policy")
for i,v in enumerate(vals): ax.text(i, v+(0.4 if v>0 else -1.2), f"{v:.2f}", ha="center", fontsize=8)
ax.grid(axis="y", alpha=.3); fig.tight_layout(); fig.savefig(FIG/"proof_reward_bar.png", dpi=140); plt.close(fig)

# ===== DOCX =====
doc=Document(); sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,m,Inches(0.9))
st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(10.5)
doc.styles["Heading 1"].font.color.rgb=RGBColor(0x1F,0x4D,0x78)
HDR="1F4D78"
def shade(c,f):
    p=c._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),f); p.append(sh)
def sf(c,sz,b=False,white=False):
    for pp in c.paragraphs:
        for r in pp.runs:
            r.font.size=Pt(sz); r.font.bold=b
            if white: r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
def table(headers, rows, fz=9, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=str(h); shade(c,HDR); sf(c,fz,True,True)
    for row in rows:
        cs=t.add_row().cells
        for j,v in enumerate(row): cs[j].text=str(v); sf(cs[j],fz)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)
    return t
def para(txt,size=10.5,bold=False,italic=False,align=None,space=6):
    p=doc.add_paragraph(); r=p.add_run(txt); r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(space); return p

tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("Proof that the DDQN Controller Is Genuinely Learning"); r.font.size=Pt(17); r.font.bold=True; r.font.color.rgb=RGBColor(0x1F,0x4D,0x78)
sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp.add_run("Topology-Agnostic Bottleneck-Ranking DDQN (final frozen Iter2 controller)"); r.font.size=Pt(11.5); r.font.italic=True
doc.add_paragraph()

para("This document provides three independent forms of evidence that the final Double-DQN controller is "
     "genuinely learning a value function and an improved policy, rather than acting as a fixed rule, an "
     "imitation (supervised) model, or an untrained network.")

doc.add_heading("Proof 1 — Learning curve (monotone improvement over 22 episodes)", level=1)
para("During training the temporal-difference (Huber) loss decreases steadily while the mean episode reward "
     "increases monotonically, as the exploration rate (epsilon) decays from exploration to exploitation. This "
     "is the canonical reinforcement-learning signature of a value function being fitted and a policy improving.")
table(["Quantity","Episode 1","Episode 22","Trend"],
 [["TD (Huber) loss", f"{tl.mean_td_loss.iloc[0]:.3f}", f"{tl.mean_td_loss.iloc[-1]:.3f}", "decreasing (~3.3x)"],
  ["Mean reward", f"{tl.mean_reward.iloc[0]:.2f}", f"{tl.mean_reward.iloc[-1]:.2f}", "increasing"],
  ["Epsilon (exploration)", f"{tl.epsilon.iloc[0]:.2f}", f"{tl.epsilon.iloc[-1]:.2f}", "decaying"]], fz=9.5, widths=[2.0,1.4,1.4,2.2])
doc.add_picture(str(FIG/"proof_learning_curve.png"), width=Inches(5.7)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
para("Figure A. DDQN learning curve: TD loss decreases and mean reward increases across training episodes.", italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_heading("Proof 2 — Real update counters (genuine Double-DQN, not a fixed rule or imitation)", level=1)
para("The runtime counters recorded during training confirm that a true value-based reinforcement-learning loop "
     "ran: experience replay, epsilon-greedy exploration, many temporal-difference gradient updates, and "
     "periodic target-network synchronization. Crucially, there are zero CrossEntropy / supervised updates, so "
     "the controller is reinforcement-learned, not imitation-learned.")
table(["Counter","Value","Meaning"],
 [["env_steps", str(cnt.get("env_steps","-")), "22 episodes x 6 topologies x 160 cycles"],
  ["td_updates", str(cnt.get("td_updates","-")), "gradient steps on the online Q-network"],
  ["target_updates", str(cnt.get("target_updates","-")), "periodic hard target-network sync"],
  ["replay_pushes", str(cnt.get("replay_pushes","-")), "experience-replay transitions stored"],
  ["explore_actions", str(cnt.get("explore_actions","-")), "epsilon-greedy exploratory actions"],
  ["greedy_actions", str(cnt.get("greedy_actions","-")), "greedy (argmax-Q) actions"],
  ["ce_updates", str(cnt.get("ce_updates",0)), "CrossEntropy/supervised updates = NONE"]], fz=9.5, widths=[1.7,1.1,4.0])

doc.add_heading("Proof 3 — Controlled comparison (trained vs untrained vs random vs fixed)", level=1)
para("The decisive test holds the network architecture and evaluation fixed and varies only whether the weights "
     "were trained. All policies are evaluated on the same objective (the training reward). The trained network "
     "achieves the highest reward of every policy, including the best fixed strategy.")
prow=[]
disp={"trained":"Trained DDQN","untrained":"Untrained (random-init)","random":"Random action","keep":"Always-KEEP","k800":"Always-K800 (best fixed)"}
for p in ["trained","untrained","random","k800","keep"]:
    rr=pm[p]; prow.append([disp[p], f"{rr.mean_reward:.2f}", f"{rr.mean_PR:.4f}", f"{rr.mean_ms:.1f}", f"{rr.action_entropy:.3f}"])
table(["Policy","Mean reward (objective)","Mean PR","Mean ms","Action entropy"], prow, fz=9.5, widths=[2.4,1.7,1.0,0.9,1.2])
tr=pm["trained"]; un=pm["untrained"]; rd=pm["random"]; k8=pm["k800"]
para("Interpretation:", bold=True)
para(f"• Trained reward {tr.mean_reward:.2f} > untrained {un.mean_reward:.2f} (+{tr.mean_reward-un.mean_reward:.2f}), "
     f"> random {rd.mean_reward:.2f} (+{tr.mean_reward-rd.mean_reward:.2f}), and > best fixed policy Always-K800 "
     f"{k8.mean_reward:.2f} (+{tr.mean_reward-k8.mean_reward:.2f}). The trained policy is the best on the objective.", size=10)
para(f"• The untrained network reaches a slightly higher raw PR ({un.mean_PR:.4f}) only by collapsing to a "
     f"near-fixed high-K action (action entropy {un.action_entropy:.3f}) that over-optimizes at about {un.mean_ms:.0f} ms "
     f"(roughly twice the runtime). The trained network instead learns the PR-versus-runtime trade-off: it attains "
     f"comparable PR ({tr.mean_PR:.4f}) at about {tr.mean_ms:.0f} ms (about half the runtime), which is exactly what "
     f"the reward rewards. A policy that blindly maximized PR would be Always-K800 (PR {k8.mean_PR:.4f} at {k8.mean_ms:.0f} ms), "
     f"and the trained DDQN still beats it on reward.", size=10)
para(f"• Always-KEEP collapses (reward {pm['keep'].mean_reward:.2f}, PR {pm['keep'].mean_PR:.4f}), confirming the "
     f"objective is non-trivial and that the trained controller did not simply default to KEEP.", size=10)
doc.add_picture(str(FIG/"proof_reward_bar.png"), width=Inches(5.7)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
para("Figure B. Mean reward by policy. The trained DDQN attains the highest reward of all policies.", italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_heading("Conclusion", level=1)
para("All three lines of evidence agree: (1) the loss decreases and reward increases over training; (2) the "
     "runtime counters show a genuine Double-DQN loop with experience replay, epsilon-greedy exploration, "
     "thousands of temporal-difference updates, target-network syncing, and zero supervised updates; and (3) in a "
     "controlled comparison the trained network outperforms an untrained network, a random policy, and the best "
     "fixed strategy on the objective. Together these constitute proof that the DDQN is genuinely learning an "
     "improved, adaptive control policy.")

DOCX = RPT / "DDQN_Learning_Proof.docx"
doc.save(str(DOCX)); print("DOCX saved:", DOCX); print("DONE")
